from comment_agent.export.documents import DocumentExporter
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def _reviews():
    reviews = {
        "2024Q3": {
            "key_variation": (
                "### Overview\nPnL movements require validation.\n\n"
                "### Key Metric Topic 1: Rates-risk PnL volatility\n"
                "**Variations:**\n- Material rates movement\n\n"
                "## Sources\n- [C1] (2024-08-15)\n- [C1] (2024-08-15)"
            ),
            "recurrent": (
                "### Overview\nRepeated data-timing issue.\n\n"
                "### Recurrent Topic 1: Valuation data timing\n"
                "**Explanations:**\n- Contextual Explanation: Delay\n\n"
                "### Technical Issues\n- Booking feed delay\n"
                "## Sources\n- [C2] (2024-09-30)"
            ),
        }
    }
    reviews["2024Q4"] = reviews["2024Q3"]
    return reviews


def _context():
    return {
        "report_id": "CAR-20260722T100000Z",
        "generated_at": "22 Jul 2026, 10:00 UTC",
        "review_status": "AI-generated — auditor validation required",
        "desks": "EQD",
        "date_range": "2024-07-01 to 2024-09-30",
        "evidence_rows": 12,
        "input_files": "cert.csv, ia.csv, pnl.csv",
    }


def _document_text(doc):
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    cells = [paragraph.text for table in doc.tables for row in table.rows
             for cell in row.cells for paragraph in cell.paragraphs]
    return "\n".join(paragraphs + cells)


def test_markdown_buffer_nonempty():
    exporter = DocumentExporter()
    buf = exporter.get_word_doc_buffer_from_markdown("# Title\n- bullet\n**bold** text")
    data = buf.getvalue()
    assert data[:2] == b"PK"  # docx is a zip
    assert len(data) > 0


def test_save_executive_summary_writes_file(tmp_path):
    exporter = DocumentExporter()
    path = exporter.save_executive_summary("## Summary\ncontent", "PnL Comment",
                                           output_dir=str(tmp_path))
    assert path.endswith(".docx")


def test_executive_review_uses_the_requested_lean_layout():
    doc = DocumentExporter().generate_executive_summary_word(
        "# Executive Summary:\nReview conclusion.", "PnL Comment",
        reviews=_reviews(), report_context=_context(),
    )
    text = _document_text(doc)

    assert "Executive Review | PnL Comment" in text
    assert "Market Activities Comment Review" not in text
    assert "CAR-20260722T100000Z" not in text
    assert "Findings Dashboard" not in text
    assert "Auditor Review and Sign-off" not in text
    assert "AI-generated" not in text
    assert "Executive Review | PnL Comment" in doc.sections[0].header.paragraphs[0].text
    assert doc.sections[0].top_margin == Inches(1)
    assert doc.styles["Report Title"].font.color.rgb == RGBColor(46, 116, 181)
    assert doc.styles["Report Title"].font.size == Pt(18)
    assert doc.styles["Normal"].font.name == "Georgia"
    for style_name in ("Normal", "List Bullet", "Heading 1", "Heading 2", "Heading 3",
                       "Report Title", "Report Subtitle", "Report Meta", "Report Callout",
                       "Source Citation"):
        rfonts = doc.styles[style_name].element.rPr.rFonts
        assert all(rfonts.get(qn(font_slot)) == "Georgia"
                   for font_slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"))
        assert all(rfonts.get(qn(theme_attribute)) is None
                   for theme_attribute in ("w:asciiTheme", "w:hAnsiTheme",
                                           "w:eastAsiaTheme", "w:cstheme"))
    assert doc.styles["Normal"].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert doc.styles["Heading 1"].font.color.rgb == RGBColor(46, 116, 181)
    table_properties = doc.tables[0]._tbl.tblPr
    assert table_properties.find(qn("w:tblW")).get(qn("w:w")) == "9360"
    assert table_properties.find(qn("w:tblInd")).get(qn("w:w")) == "120"
    footer = doc.sections[0].footer.paragraphs[0]
    assert footer.text == "Page 1"
    assert "Confidential" not in footer.text
    footer_rfonts = footer.runs[0]._element.rPr.rFonts
    assert all(footer_rfonts.get(qn(font_slot)) == "Georgia"
               for font_slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"))


def test_detailed_review_contains_quarterly_evidence_and_all_technical_issues():
    doc = DocumentExporter().generate_detailed_review_word(
        "# Executive Summary for PnL Comment", "PnL Comment", reviews=_reviews(),
        executive_summary=(
            "# Executive Summary:\nReview conclusion.\n\n"
            "## Quarterly observations\n### Quarter 2024Q3\n- Duplicate heading"
        ),
        report_context=_context(),
    )
    text = _document_text(doc)

    assert "Detailed Full Review | PnL Comment" in text
    assert "Detailed Evidence Report" not in text
    assert "Quarterly Detailed Review" in text
    assert "Quarter 2024Q3" in text
    assert "Quarter 2024Q4" in text
    assert text.count("Quarter 2024Q3") == 1
    assert "[C1] (2024-08-15)" in text
    assert "Booking feed delay" in text
    assert "Scope and Methodology" not in text
    assert "Findings Dashboard" not in text
    assert "Auditor Review and Sign-off" not in text
    source_positions = [
        index for index, paragraph in enumerate(doc.paragraphs)
        if paragraph.text == "Sources"
    ]
    assert len(source_positions) == 4
    assert text.count("[C1] (2024-08-15)") == 2

    key_positions = [
        index for index, paragraph in enumerate(doc.paragraphs)
        if paragraph.text == "Key Metric Variations"
    ]
    recurrent_positions = [
        index for index, paragraph in enumerate(doc.paragraphs)
        if paragraph.text == "Recurrent Topics"
    ]
    assert key_positions[0] < source_positions[0] < recurrent_positions[0]
    assert recurrent_positions[0] < source_positions[1]
    assert key_positions[1] < source_positions[2] < recurrent_positions[1]
    assert recurrent_positions[1] < source_positions[3]
    quarter_positions = [
        index for index, paragraph in enumerate(doc.paragraphs)
        if paragraph.text.startswith("Quarter 2024Q")
    ]
    assert doc.paragraphs[quarter_positions[1] - 1].text == ""
