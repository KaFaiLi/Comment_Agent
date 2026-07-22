"""Professional Word report generation for comment-review outputs."""

import re
from io import BytesIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

from comment_agent.logging_config import get_logger

logger = get_logger(__name__)


# ``standard_business_brief`` design tokens.  These live here so all generated
# documents use one deliberate style system rather than Word's defaults.
_CONTENT_WIDTH_DXA = 9360
_TABLE_INDENT_DXA = 120
_CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
_COLORS = {
    "ink": "000000",
    "title": "2E74B5",
    "heading": "2E74B5",
    "heading_dark": "2E74B5",
    "muted": "5B6573",
    "header_fill": "F2F4F7",
    "callout_fill": "F4F6F9",
    "border": "D5DDE5",
    "white": "FFFFFF",
}
class DocumentExporter:
    """Create executive and detailed, audit-ready Word review reports.

    The public methods retain the previous API and accept optional review and
    run-context data.  This lets existing callers continue to work while the
    Streamlit workflow can supply richer document-control information.
    """

    @staticmethod
    def _set_run_font(run, *, size=None, color=None, bold=None, italic=None):
        run.font.name = "Georgia"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
        if size is not None:
            run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = RGBColor.from_string(color)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic

    @classmethod
    def _configure_document(cls, doc: Document, comment_type: str, document_label: str):
        """Apply the standard-business-brief page, type, and furniture tokens."""
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        styles = doc.styles
        normal = styles["Normal"]
        normal.font.name = "Georgia"
        normal._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
        normal.font.size = Pt(11)
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.1

        cls._configure_style(styles["Heading 1"], 16, _COLORS["heading"], True, 16, 8,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT)
        cls._configure_style(styles["Heading 2"], 13, _COLORS["heading"], True, 12, 6,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT)
        cls._configure_style(styles["Heading 3"], 12, _COLORS["heading_dark"], True, 8, 4,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT)
        cls._ensure_style(doc, "Report Title", 18, _COLORS["title"], True, 0, 4,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)
        cls._ensure_style(doc, "Report Subtitle", 13, _COLORS["title"], False, 0, 14,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)
        cls._ensure_style(doc, "Report Meta", 9, _COLORS["muted"], False, 0, 2,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        cls._ensure_style(doc, "Report Callout", 9, _COLORS["ink"], True, 0, 0,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        cls._ensure_style(doc, "Source Citation", 9, _COLORS["muted"], False, 4, 4,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

        bullet = styles["List Bullet"]
        bullet.base_style = normal
        bullet.paragraph_format.left_indent = Inches(0.5)
        bullet.paragraph_format.first_line_indent = Inches(-0.25)
        bullet.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bullet.paragraph_format.space_after = Pt(4)
        bullet.paragraph_format.line_spacing = 1.167

        cls._configure_header_footer(section, f"{document_label} | {comment_type}")

    @classmethod
    def _configure_style(cls, style, size, color, bold, before, after, *, alignment):
        style.font.name = "Georgia"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1
        style.paragraph_format.alignment = alignment

    @classmethod
    def _ensure_style(cls, doc, name, size, color, bold, before, after, *, alignment):
        styles = doc.styles
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.base_style = styles["Normal"]
        cls._configure_style(style, size, color, bold, before, after, alignment=alignment)

    @classmethod
    def _configure_header_footer(cls, section, document_title):
        header = section.header
        paragraph = header.paragraphs[0]
        cls._clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(document_title)
        cls._set_run_font(run, size=8, color=_COLORS["muted"], bold=True)

        footer = section.footer
        paragraph = footer.paragraphs[0]
        cls._clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run("Confidential — Internal Use Only  |  Page ")
        cls._set_run_font(run, size=8, color=_COLORS["muted"])
        cls._add_page_number_field(paragraph)

    @staticmethod
    def _clear_paragraph(paragraph):
        for child in list(paragraph._p):
            paragraph._p.remove(child)

    @staticmethod
    def _add_page_number_field(paragraph):
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        value = OxmlElement("w:t")
        value.text = "1"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend([begin, instruction, separate, value, end])

    @classmethod
    def _new_document(cls, comment_type: str, document_label: str) -> Document:
        doc = Document()
        cls._configure_document(doc, comment_type, document_label)
        return doc

    @classmethod
    def _add_masthead(cls, doc, comment_type: str, document_label: str):
        title = doc.add_paragraph(style="Report Title")
        cls._set_run_font(title.add_run(f"{document_label} | {comment_type}"), size=18,
                          color=_COLORS["title"], bold=True)

    @staticmethod
    def _context_value(report_context, key, default="Not available"):
        value = (report_context or {}).get(key, default)
        if value in (None, ""):
            return default
        return str(value)

    @classmethod
    def _add_document_control(cls, doc, comment_type, report_context):
        fields = [
            ("Generated", cls._context_value(report_context, "generated_at")),
            ("Comment type", comment_type),
            ("Scope", cls._context_value(report_context, "desks", "Not specified")),
            ("Review period", cls._context_value(report_context, "date_range")),
        ]
        table = doc.add_table(rows=0, cols=4)
        table.style = "Table Grid"
        for left, right in zip(fields[::2], fields[1::2]):
            cells = table.add_row().cells
            # A four-column grid gives each label/value pair a comfortable width.
            cls._split_label_value_cells(cells, left, right)
        cls._set_table_geometry(table, [1500, 3180, 1500, 3180])
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cls._add_spacer(doc, 6)

    @classmethod
    def _split_label_value_cells(cls, cells, left, right):
        """Populate a four-column document-control row with real label/value cells."""
        cells[0].text = ""
        cells[1].text = ""
        cells[2].text = ""
        cells[3].text = ""
        cls._set_label_cell(cells[0], left[0])
        cls._set_value_cell(cells[1], left[1])
        cls._set_label_cell(cells[2], right[0])
        cls._set_value_cell(cells[3], right[1])

    @classmethod
    def _set_label_cell(cls, cell, label):
        cell.text = ""
        cls._set_cell_shading(cell, _COLORS["header_fill"])
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cls._set_run_font(paragraph.add_run(f"{label}:"), size=8,
                          color=_COLORS["muted"], bold=True)

    @classmethod
    def _set_value_cell(cls, cell, value):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        cls._set_run_font(paragraph.add_run(str(value)), size=9, color=_COLORS["ink"])

    @classmethod
    def _add_scope_and_methodology(cls, doc, report_context):
        doc.add_heading("Scope and Methodology", level=1)
        scope = cls._context_value(report_context, "desks", "Not specified")
        period = cls._context_value(report_context, "date_range")
        count = cls._context_value(report_context, "evidence_rows", "0")
        paragraph = doc.add_paragraph()
        cls._add_bold_label(paragraph, "Scope: ", f"{scope}. ")
        cls._add_bold_label(paragraph, "Review period: ", f"{period}. ")
        cls._add_bold_label(paragraph, "Evidence reviewed: ", f"{count} source record(s).")
        doc.add_paragraph(
            "The review groups alert-driven risk comments by quarter and comment type. "
            "It supports auditor assessment and does not replace professional judgement. "
            "A lack of daily comments is not, by itself, evidence of an exception."
        )

    @classmethod
    def _add_bold_label(cls, paragraph, label, value):
        cls._set_run_font(paragraph.add_run(label), size=11, color=_COLORS["ink"], bold=True)
        cls._set_run_font(paragraph.add_run(value), size=11, color=_COLORS["ink"])

    @classmethod
    def _add_findings_dashboard(cls, doc, reviews):
        rows = cls._extract_findings(reviews)
        if not rows:
            return
        doc.add_heading("Findings Dashboard", level=1)
        doc.add_paragraph(
            "Topics below are AI-generated candidates for auditor validation; no severity rating "
            "is assigned until the auditor completes their assessment."
        )
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ("ID", "Quarter", "Category", "Proposed finding", "Status")
        for cell, header in zip(table.rows[0].cells, headers):
            cls._set_header_cell(cell, header)
        for finding in rows:
            cells = table.add_row().cells
            for cell, value in zip(cells, finding):
                cell.text = ""
                paragraph = cell.paragraphs[0]
                cls._set_run_font(paragraph.add_run(value), size=9, color=_COLORS["ink"])
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cls._set_table_geometry(table, [700, 1100, 1450, 4300, 1810])
        cls._add_spacer(doc, 6)

    @staticmethod
    def _extract_findings(reviews):
        findings = []
        for quarter, review in (reviews or {}).items():
            for category, key, marker in (
                ("Key variation", "key_variation", r"^### Key Metric Topic \d+:\s*(.+)$"),
                ("Recurrent topic", "recurrent", r"^### Recurrent Topic \d+:\s*(.+)$"),
            ):
                for match in re.finditer(marker, str((review or {}).get(key, "")), re.MULTILINE):
                    findings.append((
                        f"F-{len(findings) + 1:03d}", str(quarter), category,
                        match.group(1).strip(), "Validation required",
                    ))
        return findings

    @classmethod
    def _set_header_cell(cls, cell, text):
        cell.text = ""
        cls._set_cell_shading(cell, _COLORS["header_fill"])
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cls._set_run_font(paragraph.add_run(text), size=8, color=_COLORS["ink"], bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    @classmethod
    def _add_signoff_section(cls, doc):
        doc.add_heading("Auditor Review and Sign-off", level=1)
        doc.add_paragraph(
            "Record the auditor's assessment before distributing or relying on this report."
        )
        table = doc.add_table(rows=2, cols=4)
        table.style = "Table Grid"
        for cell, header in zip(table.rows[0].cells, ("Decision", "Reviewer", "Review date", "Comments")):
            cls._set_header_cell(cell, header)
        for cell in table.rows[1].cells:
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            cls._set_run_font(paragraph.add_run(" "), size=10, color=_COLORS["ink"])
            cell.add_paragraph()
        cls._set_table_geometry(table, [1900, 2200, 1500, 3760])

    @classmethod
    def _render_markdown(cls, doc, text, *, skip_first_title=False):
        """Render the project's limited markdown dialect using Word-native styles."""
        source_mode = False
        first_title_skipped = False
        for raw_line in str(text or "").splitlines():
            line = raw_line.strip()
            if not line:
                continue
            heading = re.match(r"^(#{1,3})\s+(.+)$", line)
            if heading:
                title = heading.group(2).strip().rstrip(":")
                if skip_first_title and not first_title_skipped:
                    first_title_skipped = True
                    continue
                source_mode = title.lower() == "sources"
                level = min(len(heading.group(1)), 3)
                doc.add_heading(title, level=level)
                continue
            if line.startswith("- "):
                if source_mode:
                    paragraph = doc.add_paragraph(style="Source Citation")
                    cls._add_rich_text(paragraph, line[2:], size=9)
                else:
                    paragraph = doc.add_paragraph(style="List Bullet")
                    cls._add_rich_text(paragraph, line[2:], size=11)
                continue
            paragraph = doc.add_paragraph()
            cls._add_rich_text(paragraph, line, size=11)

    @staticmethod
    def _split_sources(markdown):
        """Split a formatted review into its analysis and source appendix."""
        body, sources = [], []
        in_sources = False
        for line in str(markdown or "").splitlines():
            if re.match(r"^#{1,3}\s+Sources\s*$", line.strip(), flags=re.IGNORECASE):
                in_sources = True
                continue
            (sources if in_sources else body).append(line)
        return "\n".join(body), sources

    @classmethod
    def _add_sources(cls, doc, source_lines):
        """Render one deduplicated citation appendix for a review section."""
        items, seen = [], set()
        for raw_line in source_lines:
            item = raw_line.strip()
            if not item:
                continue
            if item.startswith("- "):
                item = item[2:]
            key = " ".join(item.split()).casefold()
            if key and key not in seen:
                seen.add(key)
                items.append(item)
        if not items:
            return
        doc.add_heading("Sources", level=3)
        for item in items:
            paragraph = doc.add_paragraph(style="Source Citation")
            cls._add_rich_text(paragraph, item, size=9)

    @staticmethod
    def _executive_conclusion_text(text):
        """Keep only the lead executive conclusion in a detailed report.

        The detailed report owns the quarter-by-quarter sections, so rendering
        the executive summary's own quarterly digest would duplicate headings.
        """
        lines = str(text or "").splitlines()
        conclusion = []
        for raw_line in lines:
            line = raw_line.strip()
            if line.startswith("# "):
                continue
            if line.startswith("## "):
                if conclusion:
                    break
                if "executive" in line.lower():
                    continue
                break
            if line.startswith("### ") and conclusion:
                break
            if line:
                conclusion.append(line)
        return "\n".join(conclusion)

    @classmethod
    def _add_rich_text(cls, paragraph, text, *, size):
        for token in re.split(r"(\*\*.*?\*\*)", text):
            if not token:
                continue
            is_bold = token.startswith("**") and token.endswith("**")
            run = paragraph.add_run(token[2:-2] if is_bold else token)
            cls._set_run_font(run, size=size, color=_COLORS["ink"], bold=is_bold or None)

    @classmethod
    def _set_table_geometry(cls, table, widths_dxa):
        """Apply fixed, matching OOXML table and cell geometry in DXA."""
        if sum(widths_dxa) != _CONTENT_WIDTH_DXA:
            raise ValueError("Table column widths must equal the document content width")
        table.autofit = False
        tbl_pr = table._tbl.tblPr
        cls._set_xml_attribute(tbl_pr, "w:tblW", "w:w", str(_CONTENT_WIDTH_DXA))
        cls._set_xml_attribute(tbl_pr, "w:tblW", "w:type", "dxa")
        cls._set_xml_attribute(tbl_pr, "w:tblInd", "w:w", str(_TABLE_INDENT_DXA))
        cls._set_xml_attribute(tbl_pr, "w:tblInd", "w:type", "dxa")
        cls._set_xml_attribute(tbl_pr, "w:tblLayout", "w:type", "fixed")
        cls._set_table_borders(tbl_pr)

        for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, widths_dxa):
            grid_col.set(qn("w:w"), str(width))
        for row in table.rows:
            for cell, width in zip(row.cells, widths_dxa):
                cell.width = Twips(width)
                tc_pr = cell._tc.get_or_add_tcPr()
                cls._set_xml_attribute(tc_pr, "w:tcW", "w:w", str(width))
                cls._set_xml_attribute(tc_pr, "w:tcW", "w:type", "dxa")
                cls._set_cell_margins(tc_pr)

    @staticmethod
    def _set_xml_attribute(parent, tag, attribute, value):
        element = parent.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            parent.append(element)
        element.set(qn(attribute), value)

    @classmethod
    def _set_cell_margins(cls, tc_pr):
        margins = tc_pr.find(qn("w:tcMar"))
        if margins is None:
            margins = OxmlElement("w:tcMar")
            tc_pr.append(margins)
        for side, value in _CELL_MARGINS_DXA.items():
            element = margins.find(qn(f"w:{side}"))
            if element is None:
                element = OxmlElement(f"w:{side}")
                margins.append(element)
            element.set(qn("w:w"), str(value))
            element.set(qn("w:type"), "dxa")

    @staticmethod
    def _set_table_borders(tbl_pr):
        borders = tbl_pr.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = borders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), _COLORS["border"])

    @staticmethod
    def _set_cell_shading(cell, fill):
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)
        shading.set(qn("w:fill"), fill)
        shading.set(qn("w:val"), "clear")

    @staticmethod
    def _add_spacer(doc, points):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(points)
        paragraph.paragraph_format.space_before = Pt(0)

    @classmethod
    def generate_executive_summary_word(cls, executive_summary: str, comment_type: str,
                                        reviews=None, report_context=None) -> Document:
        """Build the concise management-facing executive review document."""
        doc = cls._new_document(comment_type, "Executive Review")
        cls._add_masthead(doc, comment_type, "Executive Review")
        cls._add_document_control(doc, comment_type, report_context)
        doc.add_heading("Executive Conclusion", level=1)
        cls._render_markdown(doc, executive_summary, skip_first_title=True)
        return doc

    @classmethod
    def generate_detailed_review_word(cls, markdown_content: str, comment_type: str,
                                      reviews=None, executive_summary=None,
                                      report_context=None) -> Document:
        """Build the evidence-oriented detailed report with quarterly sections."""
        doc = cls._new_document(comment_type, "Detailed Evidence Report")
        cls._add_masthead(doc, comment_type, "Detailed Evidence Report")
        cls._add_document_control(doc, comment_type, report_context)
        doc.add_heading("Executive Conclusion", level=1)
        conclusion = cls._executive_conclusion_text(executive_summary or markdown_content)
        if conclusion:
            cls._render_markdown(doc, conclusion)
        doc.add_heading("Quarterly Detailed Review", level=1)
        if reviews:
            for index, (quarter, review) in enumerate(reviews.items()):
                if index:
                    cls._add_spacer(doc, 14)
                doc.add_heading(f"Quarter {quarter}", level=2)
                key_variation, key_sources = cls._split_sources(review.get("key_variation", ""))
                recurrent, recurrent_sources = cls._split_sources(review.get("recurrent", ""))
                doc.add_heading("Key Metric Variations", level=3)
                cls._render_markdown(doc, key_variation)
                cls._add_sources(doc, key_sources)
                doc.add_heading("Recurrent Topics", level=3)
                cls._render_markdown(doc, recurrent)
                cls._add_sources(doc, recurrent_sources)
                cls._add_spacer(doc, 14)
        else:
            cls._render_markdown(doc, markdown_content, skip_first_title=True)
        return doc

    @classmethod
    def convert_markdown_to_word(cls, markdown_content: str,
                                 comment_type: str = "Comment Review") -> Document:
        """Backward-compatible conversion to the polished detailed-report format."""
        return cls.generate_detailed_review_word(markdown_content, comment_type)

    @staticmethod
    def save_word_doc(doc: Document, file_path: str):
        """Persist a Word document to the requested path."""
        doc.save(file_path)
        logger.debug("Saved Word document | %s", file_path)

    @staticmethod
    def get_word_download_buffer(doc: Document) -> BytesIO:
        """Return a Word document as a direct-download buffer."""
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def save_executive_summary(self, executive_summary: str, comment_type: str,
                               output_dir: str = "Outputs", *, reviews=None,
                               report_context=None) -> str:
        doc = self.generate_executive_summary_word(
            executive_summary, comment_type, reviews=reviews, report_context=report_context,
        )
        safe_comment_type = comment_type.replace(" ", "_").lower()
        file_path = f"{output_dir}/executive_summary_{safe_comment_type}.docx"
        self.save_word_doc(doc, file_path)
        return file_path

    def convert_and_save_markdown(self, markdown_content: str, comment_type: str,
                                  output_dir: str = "Outputs", *, reviews=None,
                                  executive_summary=None, report_context=None) -> str:
        doc = self.generate_detailed_review_word(
            markdown_content, comment_type, reviews=reviews,
            executive_summary=executive_summary, report_context=report_context,
        )
        safe_comment_type = comment_type.replace(" ", "_").lower()
        file_path = f"{output_dir}/{safe_comment_type}_full_review.docx"
        self.save_word_doc(doc, file_path)
        return file_path

    def get_word_doc_buffer_from_markdown(self, markdown_content: str,
                                          comment_type: str = "Comment Review", *,
                                          reviews=None, executive_summary=None,
                                          report_context=None) -> BytesIO:
        doc = self.generate_detailed_review_word(
            markdown_content, comment_type, reviews=reviews,
            executive_summary=executive_summary, report_context=report_context,
        )
        return self.get_word_download_buffer(doc)

    def get_word_doc_buffer_from_executive_summary(self, executive_summary: str,
                                                   comment_type: str, *, reviews=None,
                                                   report_context=None) -> BytesIO:
        doc = self.generate_executive_summary_word(
            executive_summary, comment_type, reviews=reviews, report_context=report_context,
        )
        return self.get_word_download_buffer(doc)
